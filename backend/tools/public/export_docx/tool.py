"""生成 DOCX 并上传 OSS，供前端下载。"""

from __future__ import annotations

import io
import logging
import re
import uuid
from datetime import datetime

from langchain_core.tools import tool

from backend.common.tool_outcome import format_tool_user_message

logger = logging.getLogger(__name__)

TOOL_NAME = "export_docx"

_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_UNSAFE_NAME = re.compile(r'[<>:"/\\|?*\x00-\x1f]')
# 结构标题：#～####，或 [标题]/[一级]/[二级]/[三级]（导出时转为 Word 样式，不留标记）
_HEADING_HASH_RE = re.compile(r"^(?P<marks>#{1,4})\s+(?P<text>.+?)\s*$")
_HEADING_CN_RE = re.compile(
    r"^(?:\[(?P<tag>标题|一级|二级|三级)\]|［(?P<tag2>标题|一级|二级|三级)］)\s*"
    r"(?P<text>.+?)\s*$"
)
_MD_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_MD_ITALIC_RE = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")
_MD_LINK_RE = re.compile(r"\[([^\]]+)\]\([^)]+\)")


def _safe_filename(name: str) -> str:
    name = (name or "").strip() or f"export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    name = _UNSAFE_NAME.sub("_", name)
    name = name.replace("..", "_").strip(" .")
    if not name:
        name = f"export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    if not name.lower().endswith(".docx"):
        # 去掉错误后缀再统一为 .docx
        if "." in name:
            name = name.rsplit(".", 1)[0] or name
        name = f"{name}.docx"
    return name


def _strip_inline_md(text: str) -> str:
    s = _MD_LINK_RE.sub(r"\1", text or "")
    s = _MD_BOLD_RE.sub(r"\1", s)
    s = _MD_ITALIC_RE.sub(r"\1", s)
    return s.strip()


def _parse_heading_line(line: str) -> tuple[int, str] | None:
    """返回 (Word heading level 0-3, 标题文本)；0=文档主标题。"""
    s = (line or "").strip()
    if not s:
        return None
    m = _HEADING_HASH_RE.match(s)
    if m:
        # # 主标题 → ## 一级 → ### 二级 → #### 三级
        n = len(m.group("marks"))
        return (max(0, n - 1), _strip_inline_md(m.group("text")))
    m = _HEADING_CN_RE.match(s)
    if m:
        tag = m.group("tag") or m.group("tag2") or ""
        text = _strip_inline_md(m.group("text"))
        mapping = {"标题": 0, "一级": 1, "二级": 2, "三级": 3}
        return (mapping.get(tag, 1), text)
    return None


def _set_run_font(run, *, size_pt: float, bold: bool = False) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt

    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = "Times New Roman"
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "宋体")


def _style_body_paragraph(paragraph) -> None:
    from docx.shared import Cm, Pt

    pf = paragraph.paragraph_format
    pf.first_line_indent = Cm(0.74)  # 约两字符
    pf.space_after = Pt(6)
    pf.line_spacing = 1.5
    for run in paragraph.runs:
        _set_run_font(run, size_pt=12, bold=False)


def _style_heading_paragraph(paragraph, level: int) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    sizes = {0: 18, 1: 16, 2: 14, 3: 12}
    pf = paragraph.paragraph_format
    pf.first_line_indent = Pt(0)
    pf.space_before = Pt(12 if level else 6)
    pf.space_after = Pt(8)
    if level == 0:
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragraph.runs:
        _set_run_font(run, size_pt=float(sizes.get(level, 12)), bold=True)


def _add_report_heading(doc, text: str, level: int) -> None:
    """level 0=主标题；1-3=Heading。"""
    if level <= 0:
        p = doc.add_paragraph()
        run = p.add_run(text)
        _style_heading_paragraph(p, 0)
        _set_run_font(run, size_pt=18, bold=True)
        return
    # python-docx Heading 最高用到 3
    p = doc.add_heading(text, level=min(level, 3))
    _style_heading_paragraph(p, min(level, 3))


def _add_report_body(doc, text: str) -> None:
    body = _strip_inline_md(text)
    if not body:
        return
    p = doc.add_paragraph(body)
    _style_body_paragraph(p)


def _build_report_docx(doc, text: str) -> None:
    """可商用报告排版：多级标题 + 正文首行缩进。"""
    buf_lines: list[str] = []

    def flush_body() -> None:
        nonlocal buf_lines
        if not buf_lines:
            return
        _add_report_body(doc, "\n".join(buf_lines).strip())
        buf_lines = []

    for raw in text.split("\n"):
        line = raw.rstrip()
        if not line.strip():
            flush_body()
            continue
        heading = _parse_heading_line(line.strip())
        if heading is not None:
            flush_body()
            level, title = heading
            if title:
                _add_report_heading(doc, title, level)
            continue
        # 去掉残留列表符，避免 Word 里出现 markdown 痕迹
        cleaned = re.sub(r"^[-*+]\s+", "", line.strip())
        cleaned = re.sub(r"^\d+\.\s+", "", cleaned)
        buf_lines.append(cleaned)
    flush_body()


def _build_plain_docx(doc, text: str, title: str = "") -> None:
    if title:
        doc.add_heading(title, level=1)
    if not text:
        doc.add_paragraph("")
        return
    for block in re.split(r"\n{2,}", text):
        lines = block.split("\n")
        p = doc.add_paragraph(lines[0])
        for line in lines[1:]:
            p.add_run(f"\n{line}")


def _build_docx_bytes(
    content: str,
    title: str = "",
    *,
    style: str = "plain",
) -> bytes:
    from docx import Document

    doc = Document()
    text = (content or "").replace("\r\n", "\n").replace("\r", "\n").strip()
    if (style or "plain").strip().lower() == "report":
        # 报告体：标题由正文结构标记给出，不再用文件名顶一级标题
        if text:
            _build_report_docx(doc, text)
        else:
            doc.add_paragraph("")
    else:
        _build_plain_docx(doc, text, title=title)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_docx_to_oss(
    content: str,
    filename: str = "",
    *,
    style: str = "plain",
) -> dict:
    """生成 Word 并上传 OSS，返回前端可用的下载元信息。"""
    from backend.common.oss import build_file_url, put_object, sign_get_url

    text = content if content is not None else ""
    if not str(text).strip():
        raise ValueError("内容为空，请先生成文案再导出。")

    safe_name = _safe_filename(filename)
    cover = safe_name.rsplit(".", 1)[0]
    data = _build_docx_bytes(str(text), title=cover, style=style)
    object_key = f"chat-exports/{uuid.uuid4().hex[:12]}_{safe_name}"
    put_object(object_key, data, _DOCX_MIME)
    permanent = build_file_url(object_key)
    download_url = sign_get_url(object_key, expires=3600 * 24)
    logger.info(
        "Exported docx to OSS: %s (%d bytes) style=%s",
        object_key,
        len(data),
        style,
    )
    return {
        "name": safe_name,
        "url": download_url,
        "permanent_url": permanent,
        "object_key": object_key,
        "mime_type": _DOCX_MIME,
        "file_size": len(data),
    }


def format_export_result(meta: dict) -> str:
    name = meta.get("name") or "export.docx"
    return format_tool_user_message(
        f"文档「{name}」已准备好，界面会显示下载卡片。",
        ask="请用一两句友好中文告知用户可以点击下方卡片下载到本机；",
    )


@tool(TOOL_NAME)
def export_docx(content: str, filename: str = "") -> str:
    """导出 Word（.docx）。全文已就绪，或用户短确认「好/要/导出」时调用；勿为导出再检索。

    Args:
        content: 文档全文（纯文本；勿用列表/加粗等 Markdown）
        filename: 文件名，空=自动命名
    """
    try:
        meta = export_docx_to_oss(content, filename)
    except ValueError as e:
        from backend.common.errors import tool_user_error

        return tool_user_error("导出", e)
    except Exception as e:
        logger.exception("export_docx failed")
        from backend.common.errors import tool_user_error

        return tool_user_error("导出", e)
    return format_export_result(meta)
